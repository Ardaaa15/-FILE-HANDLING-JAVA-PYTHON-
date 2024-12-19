import os
import json
from docx import Document
from docx.shared import Inches
from PIL import Image

class RestaurantOrderManagement:
    def __init__(self, orders_file='orders.json'):
        self.orders_file = orders_file
        self.orders = self.load_orders()

    def load_orders(self):
        if os.path.exists(self.orders_file):
            with open(self.orders_file, 'r') as file:
                return json.load(file)
        return []

    def save_orders(self):
        with open(self.orders_file, 'w') as file:
            json.dump(self.orders, file, indent=4)

    def create_order(self, nama_pelanggan, nama_pesanan, harga, gambar_path):
        if not os.path.exists(gambar_path):
            print(f"Error: Image file {gambar_path} not found!")
            return False

        order = {
            'id': len(self.orders) + 1,
            'nama_pelanggan': nama_pelanggan,
            'nama_pesanan': nama_pesanan,
            'harga': harga,
            'gambar_path': gambar_path,
            'status': 'diproses'
        }
        
        self.orders.append(order)
        self.save_orders()
        print(f"Pesanan untuk {nama_pelanggan} berhasil dibuat!")
        return True

    def read_orders(self):
        if not self.orders:
            print("Tidak ada pesanan.")
            return

        print("\n===== DAFTAR PESANAN =====")
        for order in self.orders:
            print(f"ID: {order['id']}")
            print(f"Nama Pelanggan: {order['nama_pelanggan']}")
            print(f"Nama Pesanan: {order['nama_pesanan']}")
            print(f"Harga: Rp {order['harga']}")
            print(f"Status: {order['status']}")
            print("------------------------")

    def update_order(self, order_id, new_status):
        valid_statuses = ['diproses', 'selesai', 'batal']
        if new_status not in valid_statuses:
            print("Status tidak valid. Gunakan: diproses/selesai/batal")
            return False

        for order in self.orders:
            if order['id'] == order_id:
                order['status'] = new_status
                self.save_orders()
                print(f"Status pesanan ID {order_id} diubah menjadi {new_status}")
                return True
        
        print(f"Pesanan dengan ID {order_id} tidak ditemukan.")
        return False

    def delete_order(self, order_id):
        for order in self.orders:
            if order['id'] == order_id:
                if order['status'] == 'batal':
                    self.orders.remove(order)
                    for i, o in enumerate(self.orders, 1):
                        o['id'] = i
                    self.save_orders()
                    print(f"Pesanan ID {order_id} berhasil dihapus.")
                    return True
                else:
                    print("Hanya pesanan dengan status 'batal' yang dapat dihapus.")
                    return False
        
        print(f"Pesanan dengan ID {order_id} tidak ditemukan.")
        return False

    def search_order(self, nama_pelanggan):
        found_orders = [order for order in self.orders if nama_pelanggan.lower() in order['nama_pelanggan'].lower()]
        
        if not found_orders:
            print(f"Tidak ada pesanan untuk pelanggan {nama_pelanggan}")
            return []

        print(f"\n===== PESANAN UNTUK {nama_pelanggan.upper()} =====")
        for order in found_orders:
            print(f"ID: {order['id']}")
            print(f"Nama Pesanan: {order['nama_pesanan']}")
            print(f"Harga: Rp {order['harga']}")
            print(f"Status: {order['status']}")
            print("------------------------")
        
        return found_orders

    def export_to_docx(self):
        doc = Document()
        doc.add_heading('Daftar Pesanan Restoran', 0)

        for order in self.orders:
            doc.add_paragraph(f"ID: {order['id']}")
            doc.add_paragraph(f"Nama Pelanggan: {order['nama_pelanggan']}")
            doc.add_paragraph(f"Nama Pesanan: {order['nama_pesanan']}")
            doc.add_paragraph(f"Harga: Rp {order['harga']}")
            doc.add_paragraph(f"Status: {order['status']}")
       
            if os.path.exists(order['gambar_path']):
                try:
                    picture = doc.add_picture(order['gambar_path'])
                    picture.width = Inches(1.5)  
                    picture.height = Inches(1.5)  
                except Exception as e:
                    print(f"Gagal menambahkan gambar: {e}")
            
            doc.add_paragraph("------------------------")

        doc.save('list_pesanan.docx')
        print("Daftar pesanan berhasil diekspor ke daftar_pesanan.docx")

def main():
    order_management = RestaurantOrderManagement()

    while True:
        print("\n===== MANAJEMEN PESANAN RESTORAN =====")
        print("1. Tambah Pesanan")
        print("2. Lihat Pesanan")
        print("3. Ubah Status Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan")
        print("6. Ekspor ke Word")
        print("7. Keluar")

        pilihan = input("Masukkan pilihan (1-7): ")

        if pilihan == '1':
            nama_pelanggan = input("Masukkan nama pelanggan: ")
            nama_pesanan = input("Masukkan nama pesanan: ")
            harga = input("Masukkan harga pesanan: ")
            gambar_path = input("Masukkan path gambar pesanan: ")
            order_management.create_order(nama_pelanggan, nama_pesanan, harga, gambar_path)

        elif pilihan == '2':
            order_management.read_orders()

        elif pilihan == '3':
            order_id = int(input("Masukkan ID pesanan yang ingin diubah: "))
            new_status = input("Masukkan status baru (diproses/selesai/batal): ")
            order_management.update_order(order_id, new_status)

        elif pilihan == '4':
            order_id = int(input("Masukkan ID pesanan yang ingin dihapus: "))
            order_management.delete_order(order_id)

        elif pilihan == '5':
            nama_pelanggan = input("Masukkan nama pelanggan: ")
            order_management.search_order(nama_pelanggan)

        elif pilihan == '6':
            order_management.export_to_docx()

        elif pilihan == '7':
            print("Terima kasih. Sampai jumpa!")
            break

        else:
            print("Pilihan tidak valid. Silakan coba lagi.")

if __name__ == "__main__":
    main()
